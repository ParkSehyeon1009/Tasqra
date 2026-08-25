import zipfile
from io import BytesIO
from unittest.mock import MagicMock

from docx import Document
from docx.shared import Inches
from PIL import Image

from app.extractors.docx_extractor import DocxExtractor
from app.extractors.embedded_image_cache import EmbeddedImageOcrCache
from app.extractors.hwpx_extractor import HwpxExtractor
from app.extractors.layout import LayoutElement


def image_bytes() -> bytes:
    buffer = BytesIO()
    Image.new("RGB", (100, 50), "white").save(buffer, format="PNG")
    return buffer.getvalue()


def ocr_element() -> LayoutElement:
    return LayoutElement(
        x=10,
        y=10,
        x2=80,
        y2=30,
        content="반복 이미지",
        source="ocr",
        confidence=0.9,
    )


def test_cache_reuses_successful_ocr_results():
    ocr = MagicMock()
    ocr.extract.return_value = [ocr_element()]
    cache = EmbeddedImageOcrCache()
    content = image_bytes()

    first_image, first_elements = cache.extract(content, ocr)
    second_image, second_elements = cache.extract(content, ocr)

    ocr.extract.assert_called_once()
    assert cache.hit_count == 1
    assert first_image is second_image
    assert first_elements is not second_elements
    assert [item.content for item in second_elements] == ["반복 이미지"]


def test_cache_also_reuses_empty_ocr_results():
    ocr = MagicMock()
    ocr.extract.return_value = []
    cache = EmbeddedImageOcrCache()
    content = image_bytes()

    cache.extract(content, ocr)
    cache.extract(content, ocr)

    ocr.extract.assert_called_once()
    assert cache.hit_count == 1


def test_docx_repeated_image_runs_ocr_once_but_keeps_each_review_page(tmp_path):
    image_path = tmp_path / "logo.png"
    image_path.write_bytes(image_bytes())
    document_path = tmp_path / "repeated.docx"
    document = Document()
    document.add_picture(str(image_path), width=Inches(1))
    document.add_picture(str(image_path), width=Inches(1))
    document.save(document_path)
    ocr = MagicMock()
    ocr.extract.return_value = [ocr_element()]

    result = DocxExtractor(ocr).extract(str(document_path))

    ocr.extract.assert_called_once()
    assert len(result.review_pages) == 2
    assert result.content.count("반복 이미지") == 2


def test_hwpx_repeated_image_runs_ocr_once_but_keeps_each_review_page(tmp_path):
    document_path = tmp_path / "repeated.hwpx"
    section = """<sec>
      <p><run><pic><img binaryItemIDRef="image1" /></pic></run></p>
      <p><run><pic><img binaryItemIDRef="image1" /></pic></run></p>
    </sec>"""
    manifest = """<manifest>
      <item id="image1" href="BinData/image.png" media-type="image/png" />
    </manifest>"""
    with zipfile.ZipFile(document_path, "w") as archive:
        archive.writestr("Contents/section0.xml", section)
        archive.writestr("Contents/content.hpf", manifest)
        archive.writestr("BinData/image.png", image_bytes())
    ocr = MagicMock()
    ocr.extract.return_value = [ocr_element()]

    result = HwpxExtractor(ocr).extract(str(document_path))

    ocr.extract.assert_called_once()
    assert len(result.review_pages) == 2
    assert result.content.count("반복 이미지") == 2
